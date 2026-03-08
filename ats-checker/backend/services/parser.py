"""
services/parser.py
Extracts raw text from PDF and DOCX files.
Also detects basic formatting issues during extraction.
"""

import io
import re
from typing import Tuple, List


def parse_pdf(file_bytes: bytes) -> Tuple[str, List[str]]:
    """
    Extract text from PDF bytes.
    Returns (raw_text, format_warnings)
    """
    import pdfplumber

    format_warnings = []
    full_text = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            # Check for tables (ATS can't read tables)
            tables = page.extract_tables()
            if tables:
                format_warnings.append(f"table_detected_page_{page_num}")

            # Check for multiple columns heuristically
            words = page.extract_words()
            if words:
                x_positions = [w["x0"] for w in words]
                x_range = max(x_positions) - min(x_positions)
                # If words are spread in 2+ distinct clusters left/right → columns
                left_words = [x for x in x_positions if x < page.width * 0.4]
                right_words = [x for x in x_positions if x > page.width * 0.6]
                if len(left_words) > 10 and len(right_words) > 10:
                    format_warnings.append(f"multi_column_detected_page_{page_num}")

            text = page.extract_text()
            if text:
                full_text.append(text)

        # Check for images that might be header/logo
        for page in pdf.pages:
            if hasattr(page, "images") and page.images:
                format_warnings.append("images_detected")
                break

    return "\n".join(full_text), list(set(format_warnings))


def parse_docx(file_bytes: bytes) -> Tuple[str, List[str]]:
    """
    Extract text from DOCX bytes.
    Returns (raw_text, format_warnings)
    """
    from docx import Document
    from docx.oxml.ns import qn

    format_warnings = []
    doc = Document(io.BytesIO(file_bytes))
    full_text = []

    # Extract paragraph text
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    # Check for tables in doc (bad for ATS)
    if doc.tables:
        format_warnings.append("table_detected")

    # Check for text boxes (ATS can't read these)
    for shape in doc.inline_shapes:
        format_warnings.append("text_box_or_image_detected")
        break

    # Check headers/footers — if they contain lots of text it's a problem
    for section in doc.sections:
        header_text = ""
        if section.header:
            for para in section.header.paragraphs:
                header_text += para.text
        if len(header_text.strip()) > 30:
            format_warnings.append("important_content_in_header")

    return "\n".join(full_text), list(set(format_warnings))


def clean_text(text: str) -> str:
    """Normalize whitespace, remove junk characters."""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)  # remove non-ASCII
    text = text.strip()
    return text